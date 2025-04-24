import os
import json

from docx import Document


def extract_audio_info(file_path):
    """
    Extracts metadata (audio info) from a .docx file.
    """
    doc = Document(file_path)

    audio_file_name = None
    audio_file_duration = None
    published_date = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if "Audio File Name:" in text:
            audio_file_name = text.split(":")[1].strip()
        elif "Audio File Duration:" in text:
            duration_parts = text.split(":")
            if len(duration_parts) >= 3:
                audio_file_duration = ":".join(duration_parts[1:]).strip()
        elif "Published Date:" in text:
            published_date = text.split(":")[1].strip()

    return {
        "Audio File Name": audio_file_name,
        "Audio File Duration": audio_file_duration,
        "Published Date": published_date
    }


def read_word_doc(doc_path):
    """
    Reads and returns the full text of a Word document.
    """
    doc = Document(doc_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return ' '.join(full_text)


def split_text_into_chunks(text, chunk_size=1000):
    """
    Splits text into chunks of specified size.
    """
    words = text.split()
    chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]

    if len(chunks) > 1 and len(chunks[-1].split()) < 500:
        chunks[-2] += ' ' + chunks[-1]
        chunks.pop()

    return chunks


def save_chunks_to_files(chunks, output_folder):
    """
    Saves text chunks to individual files in a folder.
    """
    file_paths = []
    for i, chunk in enumerate(chunks, start=1):
        chunk_file_path = os.path.join(output_folder, f"chunk_{i}.txt")
        with open(chunk_file_path, "w", encoding="utf-8") as file:
            file.write(chunk)
        file_paths.append(chunk_file_path)
    return file_paths


def process_document_file(doc_file_path, output_folder):
    """
    Processes a Word document: extracts metadata, chunks, and saves to files.
    """
    if not os.path.exists(doc_file_path):
        raise FileNotFoundError(f"Document not found: {doc_file_path}")

    # Extract metadata
    audio_info = extract_audio_info(doc_file_path)

    # Extract text and split into chunks
    text = read_word_doc(doc_file_path)
    os.makedirs(output_folder, exist_ok=True)
    chunks = split_text_into_chunks(text, chunk_size=1000)
    chunk_files = save_chunks_to_files(chunks, output_folder)

    return {
        "Audio Metadata": audio_info,
        "Chunk Files": chunk_files
    }

def find_file_in_subfolders(filename , main_folder ):
    for root, dirs, files in os.walk(main_folder):
        if f"{filename}.docx" in files:
            return os.path.join(root, f"{filename}.docx")
    return None

def process_doc_chunks(doc_folder , json_file_path , main_folder):
     os.makedirs(doc_folder, exist_ok=True)
     process_json_and_docs(json_file_path, doc_folder , main_folder)

def process_json_and_docs(json_file, doc_folder , main_folder):
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    
    for entry in data:
        named_as = entry["Named as"]
        channel_name = entry["Channel"]
        doc_file_path = find_file_in_subfolders(named_as , main_folder)
        
        try:
            if os.path.exists(doc_file_path):
                print(f"Processing {doc_file_path}")
                audio_info = extract_audio_info(doc_file_path)
                entry['Audio Metadata'] = audio_info
                entry['chunks'] = {}


                text = read_word_doc(doc_file_path)
                output_folder = f"{doc_folder}/{channel_name}_{named_as}_Document"
                os.makedirs(output_folder, exist_ok=True)
                

                chunks = split_text_into_chunks(text, chunk_size=1000)
               
                chunks_name = save_chunks_to_files(chunks, output_folder)
                print(chunks_name)
                for i , chunk in enumerate(chunks_name):
                    entry['chunks'][f'chunk_{i}'] = {"filePath" : chunk}
                print(f"Successfully saved {len(chunks)} chunks in {output_folder}.")
            else:
                print(f"Document {doc_file_path} not found.")
        except Exception as error:
            print(f"An error occurred: {error}")
    with open(json_file, 'w', encoding='utf-8') as outfile:
        json.dump(data, outfile, indent=4, ensure_ascii=False)
    print(f"Updated metadata saved to {json_file}")
